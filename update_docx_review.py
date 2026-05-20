from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DASHBOARD = Path(__file__).resolve().parent
DATA_JS = DASHBOARD / "data.js"
DOCX_PATH = ROOT / "EBOD_Literature Review.docx"
BACKUP_PATH = ROOT / "EBOD_Literature Review.backup_before_49_update.docx"


def load_payload() -> dict:
    text = DATA_JS.read_text(encoding="utf-8")
    return json.loads(text.split("=", 1)[1].rsplit(";", 1)[0])


def clean(text: str) -> str:
    if text is None:
        return "Not reported"
    value = str(text).strip()
    if not value:
        return "Not reported"
    value = re.sub(r"\s+", " ", value)
    return value


def sentence(text: str) -> str:
    value = clean(text)
    value = value.replace(" .", ".")
    if value == "Not reported":
        return value
    if value[-1] not in ".!?":
        value += "."
    return value


def normalize(text: str) -> str:
    value = clean(text).lower()
    value = re.sub(r"[^a-z0-9]+", "", value)
    return value


def compact_list(text: str) -> str:
    value = clean(text)
    value = value.replace(" | ", "; ")
    value = value.replace("•", "-")
    return value


def split_items(text: str) -> list[str]:
    raw = str(text or "").replace("\r", "\n")
    raw = raw.replace("•", "\n• ")
    raw = raw.replace(" | ", "\n")
    lines = []
    for piece in raw.splitlines():
        part = piece.strip()
        if not part:
            continue
        part = re.sub(r"^[*\-•]\s*", "", part)
        part = re.sub(r"^\d+\.\s*", "", part)
        part = clean(part)
        if part and part.lower() != "not reported":
            lines.append(part)
    deduped = []
    seen = set()
    for item in lines:
        key = normalize(item)
        if key not in seen:
            seen.add(key)
            deduped.append(item)
    return deduped


def bullet_block(items: list[str], bullet: str = "-") -> str:
    if not items:
        return "Not reported"
    return "\n\n".join(f"{bullet} {sentence(item)}" for item in items)


def first_sentence(text: str) -> str:
    value = clean(text)
    parts = re.split(r"(?<=[.!?])\s+", value)
    return sentence(parts[0] if parts else value)


def objective_from_screening(paper: dict) -> str:
    category = clean(paper["category"])
    output = clean(paper["main_output"])
    country = clean(paper["country"])
    exposures = clean(paper["exposures"])
    outcomes = clean(paper["outcomes"])
    if output != "Not reported":
        return sentence(output)
    return sentence(
        f"This paper examines {exposures.lower()} in {country} and links them to {outcomes.lower()} using a {category.lower()} approach"
    )


def workflow_from_screening(paper: dict) -> str:
    exposures = clean(paper["exposures"])
    outcomes = clean(paper["outcomes"])
    framework = clean(paper["framework"])
    if "comparative risk assessment" in framework.lower():
        return sentence(
            f"Identify {exposures.lower()} → define exposed population and health outcomes ({outcomes.lower()}) → apply relative risk or exposure-response evidence → calculate attributable burden → interpret results within the {framework} framework"
        )
    return sentence(
        f"Identify {exposures.lower()} → link exposure to {outcomes.lower()} → apply the {framework} approach → estimate attributable burden and interpret policy relevance"
    )


def modeling_items_from_screening(paper: dict) -> list[str]:
    framework = clean(paper["framework"])
    items = split_items(framework)
    if not items:
        items = [framework]
    if clean(paper["rr"]).lower().startswith("yes"):
        items.append("Relative Risk / exposure-response integration")
    if clean(paper["af"]).lower().startswith("yes"):
        items.append("Attributable Fraction / Population Attributable Fraction calculations")
    if clean(paper["daly"]).lower().startswith("yes"):
        items.append("DALY / YLL / YLD burden estimation")
    if clean(paper["uncertainty"]).lower().startswith("yes"):
        items.append("Uncertainty or sensitivity analysis")
    return split_items("\n".join(items))


def output_items_from_screening(paper: dict) -> list[str]:
    items = split_items(paper["main_output"])
    if clean(paper["daly"]).lower().startswith("yes"):
        items.append("Generated burden outputs in deaths, DALYs, or related burden metrics")
    if clean(paper["uncertainty"]).lower().startswith("yes"):
        items.append("Reported uncertainty-aware burden interpretation")
    return split_items("\n".join(items))


def notes_from_screening(paper: dict) -> str:
    reasons = split_items(paper["reason"])
    notes = []
    notes.append(f"India relevance: {clean(paper['india_relevance'])}.")
    if reasons:
        notes.append(reasons[0])
    if clean(paper["shortlist"]).lower().startswith("yes"):
        notes.append("This paper is strong enough to be considered part of the core reusable evidence base.")
    elif clean(paper["shortlist"]).lower().startswith("maybe"):
        notes.append("This paper is useful as supporting evidence but may require more context before direct reuse.")
    if clean(paper["rr"]).lower().startswith("yes"):
        notes.append("Useful for anchoring RR or exposure-response logic.")
    if clean(paper["af"]).lower().startswith("yes"):
        notes.append("Useful for attributable fraction calculations.")
    if clean(paper["daly"]).lower().startswith("yes"):
        notes.append("Useful for DALY/YLL/YLD conversion logic.")
    if clean(paper["uncertainty"]).lower().startswith("yes"):
        notes.append("Includes uncertainty handling that can strengthen scientific defensibility.")
    return bullet_block(notes, bullet="-")


def citation_block(paper: dict) -> str:
    journal = clean(paper["journal"])
    citation = clean(paper["citation"])
    return f"{journal}\n\n{citation}"


def objective_from_deep(deep: dict, paper: dict) -> str:
    return sentence(deep["objective"] if clean(deep["objective"]) != "Not reported" else paper["main_output"])


def workflow_from_deep(deep: dict, paper: dict) -> str:
    workflow = clean(deep["workflow"])
    if workflow != "Not reported":
        return sentence(workflow)
    return workflow_from_screening(paper)


def input_items_from_deep(deep: dict, paper: dict) -> list[str]:
    items = split_items(deep["inputs"])
    if not items:
        items = split_items(f"Exposure(s): {paper['exposures']}\nHealth outcome(s): {paper['outcomes']}")
    return items


def modeling_items_from_deep(deep: dict, paper: dict) -> list[str]:
    items = split_items(deep["framework"])
    items.extend(split_items(deep["effect_measures"]))
    if not items:
        items = modeling_items_from_screening(paper)
    return split_items("\n".join(items))


def output_items_from_deep(deep: dict, paper: dict) -> list[str]:
    items = split_items(deep["quant_outputs"])
    items.extend(split_items(deep["findings"]))
    if not items:
        items = output_items_from_screening(paper)
    return split_items("\n".join(items))


def notes_from_deep(deep: dict, paper: dict) -> str:
    notes: list[str] = []
    notes.append(f"India relevance: {clean(paper['india_relevance'])}.")
    notes.append("This is one of the top 10 highlighted papers for direct methodological reuse.")
    notes.extend([f"We can reuse: {item}" for item in split_items(deep["direct_reuse"])[:6]])
    if clean(deep["india_relevance"]) != "Not reported":
        notes.append(clean(deep["india_relevance"]))
    limitation = first_sentence(deep["limitations"])
    if limitation != "Not reported.":
        notes.append(f"Limitation: {limitation}")
    return bullet_block(notes, bullet="-")


def build_top10_lookup(top10: list[dict]) -> dict[str, dict]:
    lookup = {}
    for paper in top10:
        lookup[normalize(paper["title"])] = paper
        lookup[normalize(paper["citation"])] = paper
    return lookup


def match_top10(paper: dict, lookup: dict[str, dict]) -> dict | None:
    for key in (normalize(paper["title"]), normalize(paper["citation"])):
        if key in lookup:
            return lookup[key]
    return None


def make_row_content(paper: dict, deep: dict | None) -> list[str]:
    paper_and_journal = citation_block(paper)

    if deep:
        study_objective = objective_from_deep(deep, paper)
        method_followed = workflow_from_deep(deep, paper)
        data_inputs = bullet_block(input_items_from_deep(deep, paper), bullet="-")
        modeling = bullet_block(modeling_items_from_deep(deep, paper), bullet="-")
        outputs = bullet_block(output_items_from_deep(deep, paper), bullet="-")
        notes = notes_from_deep(deep, paper)
    else:
        study_objective = objective_from_screening(paper)
        method_followed = workflow_from_screening(paper)
        data_inputs = bullet_block(
            split_items(f"{paper['exposures']}\n{paper['outcomes']}"), bullet="-"
        )
        modeling = bullet_block(modeling_items_from_screening(paper), bullet="-")
        outputs = bullet_block(output_items_from_screening(paper), bullet="-")
        notes = notes_from_screening(paper)

    return [
        paper_and_journal,
        study_objective,
        method_followed,
        data_inputs,
        modeling,
        outputs,
        notes,
    ]


def replace_table_rows(doc: Document, rows: list[list[str]]) -> None:
    table = doc.tables[0]
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            cell.text = ""
            parts = str(value).split("\n")
            if parts:
                cell.paragraphs[0].text = parts[0]
                for part in parts[1:]:
                    cell.add_paragraph(part)


def main() -> None:
    payload = load_payload()
    inventory = payload["inventory"]
    top10_lookup = build_top10_lookup(payload["top10"])
    rows = []
    for paper in inventory:
        deep = match_top10(paper, top10_lookup)
        rows.append(make_row_content(paper, deep))

    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(DOCX_PATH)
    replace_table_rows(doc, rows)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
